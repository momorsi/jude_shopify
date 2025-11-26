#!/usr/bin/env python3
"""
Generate PDF and DOC documentation from markdown file
"""

import re
import sys
from pathlib import Path

def markdown_to_docx(markdown_file, output_file):
    """Convert markdown to DOCX using python-docx"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:], level=4)
            # Code blocks
            elif line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    para = doc.add_paragraph()
                    run = para.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            # Regular paragraphs
            else:
                # Handle bold and italic
                para = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and len(part) > 1:
                        run = para.add_run(part[1:-1])
                        run.italic = True
                    else:
                        para.add_run(part)
            
            i += 1
        
        doc.save(output_file)
        print(f"✅ DOCX file created: {output_file}")
        return True
        
    except ImportError:
        print("❌ python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        return markdown_to_docx(markdown_file, output_file)
    except Exception as e:
        print(f"❌ Error creating DOCX: {str(e)}")
        return False

def markdown_to_pdf(markdown_file, output_file):
    """Convert markdown to PDF using markdown and weasyprint"""
    try:
        import markdown
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        
        # Add CSS styling
        html_document = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'Calibri', Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-top: 30px;
                }}
                h2 {{
                    color: #34495e;
                    border-bottom: 2px solid #95a5a6;
                    padding-bottom: 5px;
                    margin-top: 25px;
                }}
                h3 {{
                    color: #555;
                    margin-top: 20px;
                }}
                h4 {{
                    color: #666;
                    margin-top: 15px;
                }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                    font-size: 9pt;
                }}
                pre {{
                    background-color: #f8f8f8;
                    border: 1px solid #ddd;
                    border-radius: 5px;
                    padding: 10px;
                    overflow-x: auto;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                hr {{
                    border: none;
                    border-top: 1px solid #ddd;
                    margin: 20px 0;
                }}
                strong {{
                    color: #2c3e50;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        HTML(string=html_document).write_pdf(output_file)
        print(f"✅ PDF file created: {output_file}")
        return True
        
    except ImportError as e:
        print(f"❌ Required packages not installed: {str(e)}")
        print("Installing markdown and weasyprint...")
        import subprocess
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "markdown", "weasyprint"])
            return markdown_to_pdf(markdown_file, output_file)
        except Exception as install_error:
            print(f"❌ Installation failed: {str(install_error)}")
            print("\nAlternative: Use pandoc if available:")
            print("  pandoc PROJECT_DOCUMENTATION.md -o PROJECT_DOCUMENTATION.pdf")
            return False
    except Exception as e:
        print(f"❌ Error creating PDF: {str(e)}")
        print("\nAlternative: Use pandoc if available:")
        print("  pandoc PROJECT_DOCUMENTATION.md -o PROJECT_DOCUMENTATION.pdf")
        return False

def main():
    """Main function"""
    markdown_file = Path("PROJECT_DOCUMENTATION.md")
    
    if not markdown_file.exists():
        print(f"❌ Markdown file not found: {markdown_file}")
        return
    
    print("=" * 60)
    print("Documentation Generator")
    print("=" * 60)
    print(f"Source: {markdown_file}")
    print()
    
    # Generate DOCX
    docx_file = Path("PROJECT_DOCUMENTATION.docx")
    print("Generating DOCX file...")
    markdown_to_docx(markdown_file, docx_file)
    print()
    
    # Generate PDF
    pdf_file = Path("PROJECT_DOCUMENTATION.pdf")
    print("Generating PDF file...")
    markdown_to_pdf(markdown_file, pdf_file)
    print()
    
    print("=" * 60)
    print("Documentation generation complete!")
    print("=" * 60)

if __name__ == "__main__":
    main()

